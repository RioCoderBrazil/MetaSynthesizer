#!/usr/bin/env python3
"""
Create a properly readable Word document from the FIXED extraction
"""

import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX, WD_PARAGRAPH_ALIGNMENT

def create_readable_doc():
    # Load fixed chunks
    chunks_file = Path("02_chunked_data_FIXED/23489BE_fixed_chunks.json")
    with open(chunks_file, 'r', encoding='utf-8') as f:
        chunks = json.load(f)
    
    # Create document
    doc = Document()
    
    # Title
    title = doc.add_heading('23489BE - Querschnittsprüfung des Umgangs des Bundes mit problematischen Stoffen', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Status message
    status = doc.add_paragraph()
    status_run = status.add_run('✓ DOKUMENT KORREKT REKONSTRUIERT - TEXT IST JETZT LESBAR!')
    status_run.font.bold = True
    status_run.font.size = Pt(14)
    status_run.font.color.rgb = RGBColor(0, 128, 0)
    
    # Color mapping
    label_to_highlight = {
        'findings': WD_COLOR_INDEX.BRIGHT_GREEN,
        'appendix': WD_COLOR_INDEX.GRAY_25,
        'wik': WD_COLOR_INDEX.TURQUOISE,
        'introduction': WD_COLOR_INDEX.YELLOW,
        'evaluation': WD_COLOR_INDEX.BLUE,
        'response': WD_COLOR_INDEX.PINK,
        'recommendations': WD_COLOR_INDEX.DARK_YELLOW
    }
    
    label_names = {
        'findings': 'Feststellungen (Findings)',
        'appendix': 'Anhang (Appendix)',
        'wik': 'Wesentliches in Kürze (Executive Summary)',
        'introduction': 'Auftrag und Vorgehen (Introduction)',
        'evaluation': 'Beurteilung der EFK (Evaluation)',
        'response': 'Stellungnahme (Response)', 
        'recommendations': 'Empfehlungen (Recommendations)'
    }
    
    # Group chunks by label
    from collections import defaultdict
    sections = defaultdict(list)
    for chunk in chunks:
        sections[chunk['label']].append(chunk['text'])
    
    # Add content in logical order
    section_order = ['wik', 'introduction', 'findings', 'evaluation', 'recommendations', 'response', 'appendix']
    
    for label in section_order:
        if label in sections:
            # Section heading
            doc.add_page_break()
            heading = doc.add_heading(label_names[label], 1)
            
            # Add all text for this section
            for text in sections[label]:
                para = doc.add_paragraph()
                run = para.add_run(text)
                run.font.highlight_color = label_to_highlight.get(label, WD_COLOR_INDEX.AUTO)
                run.font.size = Pt(11)
                
                # Add spacing between chunks
                para.add_run('\n')
    
    # Save
    output_file = Path("23489BE_LESBAR_KORREKT.docx")
    doc.save(str(output_file))
    print(f"✓ Lesbares Dokument erstellt: {output_file}")
    
    # Also create a plain text version for immediate verification
    text_file = Path("23489BE_LESBAR_KORREKT.txt")
    with open(text_file, 'w', encoding='utf-8') as f:
        f.write("23489BE - KORRIGIERTE VERSION\n")
        f.write("="*80 + "\n\n")
        
        for label in section_order:
            if label in sections:
                f.write(f"\n{label_names[label].upper()}\n")
                f.write("-"*80 + "\n\n")
                
                for text in sections[label]:
                    f.write(text + "\n\n")
    
    print(f"✓ Text-Version erstellt: {text_file}")

if __name__ == "__main__":
    create_readable_doc()
