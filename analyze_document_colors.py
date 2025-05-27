#!/usr/bin/env python3
"""
Analyze document colors and extract the legend from the last page
"""

from docx import Document
import json
from collections import defaultdict
import sys


def extract_legend(docx_path):
    """Extract legend from the last page of the document"""
    doc = Document(docx_path)
    
    # Get all paragraphs
    all_paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            all_paragraphs.append(para)
    
    # Look for legend in the last paragraphs
    legend_info = []
    legend_started = False
    
    # Check last 20 paragraphs for legend
    for para in all_paragraphs[-20:]:
        text = para.text.strip()
        
        # Check if this might be a legend entry
        if any(keyword in text.lower() for keyword in ['wesentliches', 'einleitung', 'feststellung', 'beurteilung', 'empfehlung', 'stellungnahme', 'hintergrund', 'findings', 'evaluation', 'recommendations']):
            # Get the color from the runs
            color = None
            for run in para.runs:
                if run.font.highlight_color:
                    color = str(run.font.highlight_color)
                    break
            
            legend_info.append({
                'text': text,
                'color': color
            })
            legend_started = True
        elif legend_started and not text:
            # Empty line after legend started might indicate end
            break
    
    return legend_info


def analyze_colors_in_document(docx_path):
    """Analyze all colors used in the document"""
    doc = Document(docx_path)
    
    color_stats = defaultdict(list)
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            # Check for highlighting
            for run in para.runs:
                if run.font.highlight_color:
                    color = str(run.font.highlight_color)
                    color_stats[color].append({
                        'text_preview': text[:100] + '...' if len(text) > 100 else text,
                        'paragraph_index': i
                    })
                    break
    
    return color_stats


def main():
    if len(sys.argv) < 2:
        print("Usage: python analyze_document_colors.py <docx_file>")
        sys.exit(1)
    
    docx_path = sys.argv[1]
    
    print(f"Analyzing document: {docx_path}")
    print("=" * 80)
    
    # Extract legend
    print("\nExtracted Legend from last page:")
    print("-" * 40)
    legend = extract_legend(docx_path)
    for entry in legend:
        print(f"Text: {entry['text']}")
        print(f"Color: {entry['color']}")
        print("-" * 40)
    
    # Analyze all colors
    print("\nColor usage throughout document:")
    print("-" * 40)
    color_stats = analyze_colors_in_document(docx_path)
    
    for color, occurrences in color_stats.items():
        print(f"\nColor: {color}")
        print(f"Number of occurrences: {len(occurrences)}")
        print("Sample texts:")
        for i, occ in enumerate(occurrences[:3]):  # Show first 3 examples
            print(f"  {i+1}. {occ['text_preview']}")
    
    print("\n" + "=" * 80)


if __name__ == "__main__":
    main()
