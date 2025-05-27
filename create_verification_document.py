#!/usr/bin/env python3
"""
Create a verification document with chunks and colors
"""

import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX

# Color mapping
LABEL_TO_COLOR = {
    'wik': WD_COLOR_INDEX.TURQUOISE,
    'introduction': WD_COLOR_INDEX.YELLOW,
    'findings': WD_COLOR_INDEX.BRIGHT_GREEN,
    'evaluation': WD_COLOR_INDEX.BLUE,
    'recommendations': WD_COLOR_INDEX.DARK_YELLOW,
    'response': WD_COLOR_INDEX.PINK,
    'appendix': WD_COLOR_INDEX.GRAY_50
}

# Label to German description
LABEL_TO_NAME = {
    'wik': 'WIK / Zusammenfassung / Das Wichtigste in Kürze',
    'introduction': 'Einleitung / Hintergrund / Methode / Prüffragen',
    'findings': 'Feststellungen',
    'evaluation': 'Beurteilung',
    'recommendations': 'Empfehlungen',
    'response': 'Stellungnahme Geprüfte',
    'appendix': 'Anhang'
}

def create_verification_document(chunks_file, output_file):
    """Create a verification document with colored chunks"""
    # Load chunks
    with open(chunks_file, 'r', encoding='utf-8') as f:
        chunks = json.load(f)
    
    # Create a new document
    doc = Document()
    
    # Add a title
    doc.add_heading('Verification Document', 0)
    doc.add_paragraph(f'Source: {chunks_file.name}')
    doc.add_paragraph(f'Total chunks: {len(chunks)}')
    
    # Add a color legend
    p = doc.add_paragraph()
    p.add_run('COLOR LEGEND:').bold = True
    
    for label, name in LABEL_TO_NAME.items():
        p = doc.add_paragraph()
        run = p.add_run(f'{name} ({label})')
        run.font.highlight_color = LABEL_TO_COLOR[label]
    
    doc.add_paragraph()
    doc.add_heading('DOCUMENT CONTENT', 1)
    
    # Sort chunks by position (assume they have a chunk_id with numbering)
    chunks.sort(key=lambda x: x['chunk_id'])
    
    # Add each chunk with appropriate highlighting
    for i, chunk in enumerate(chunks):
        # Add section heading
        doc.add_heading(f"Chunk {i+1}: {chunk['label'].upper()} (Pages {chunk['start_page']}-{chunk['end_page']})", 2)
        
        # Add the chunk text with highlighting
        p = doc.add_paragraph()
        run = p.add_run(chunk['text'])
        run.font.highlight_color = LABEL_TO_COLOR[chunk['label']]
        
        # Add separator
        doc.add_paragraph('---------------------------------------------------')
    
    # Save the document
    doc.save(output_file)
    print(f"✅ Created verification document: {output_file}")

if __name__ == "__main__":
    # Choose a document to verify (using 23489BE as example)
    document_name = "23489BE Prüfbericht V03 - Querschnittsprüfung des Umgangs des Bundes mit problematischen Stoffen - coloured"
    
    chunks_file = Path(f"02_chunked_data_CORRECTED/{document_name}_chunks.json")
    output_file = Path(f"VERIFICATION_{document_name}.docx")
    
    create_verification_document(chunks_file, output_file)
