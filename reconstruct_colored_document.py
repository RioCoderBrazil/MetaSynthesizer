#!/usr/bin/env python3
"""
Reconstruct a Word document from labeled chunks to verify labeling was applied correctly.
Shows chunks with their associated colors and page numbers.
"""

import json
import logging
from pathlib import Path
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import argparse

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Color mapping based on the labeling schema
COLOR_MAP = {
    'wik': RGBColor(0, 255, 255),           # Cyan - Wesentliches in Kürze
    'summary': RGBColor(0, 255, 255),       # Cyan - Summary (same as wik)
    'introduction': RGBColor(255, 255, 0),   # Yellow - Auftrag und Vorgehen
    'findings': RGBColor(0, 255, 0),         # Green - Findings/Feststellungen
    'evaluation': RGBColor(0, 0, 255),       # Blue - Beurteilung/Evaluation
    'recommendations': RGBColor(255, 140, 0), # Dark Orange - Empfehlungen
    'response': RGBColor(255, 192, 203),     # Pink - Stellungnahme
    'appendix': RGBColor(128, 128, 128),     # Gray - Anhang/Annex
    'empfehlung': RGBColor(255, 140, 0),     # Dark Orange - Same as recommendations
    'background': RGBColor(135, 206, 235),   # Sky Blue - Background
    'theme': RGBColor(255, 182, 193),        # Light Pink - Themenbericht
    'core_issue': RGBColor(255, 99, 71),     # Tomato - Kernproblem
    'info_a': RGBColor(152, 251, 152),       # Pale Green - Info A
    'info_b': RGBColor(255, 255, 224),       # Light Yellow - Info B
    'info_c': RGBColor(230, 230, 250),       # Lavender - Info C
    'flank_a': RGBColor(255, 228, 225),      # Misty Rose - Flankierend A
    'flank_b': RGBColor(255, 239, 213),      # Papaya Whip - Flankierend B
    'flank_c': RGBColor(240, 248, 255),      # Alice Blue - Flankierend C
    'risk': RGBColor(255, 140, 0),           # Dark Orange - Risiken
    'cost': RGBColor(218, 165, 32),          # Goldenrod - Kosten
    'timeline': RGBColor(176, 196, 222),     # Light Steel Blue - Zeithorizont
    'misc': RGBColor(211, 211, 211),         # Light Gray - Sonstiges
    'unknown': RGBColor(128, 128, 128),      # Gray - Unknown (same as appendix)
    'default': RGBColor(255, 255, 255)       # White - No label/Unknown
}

def get_label_name(label):
    """Get human-readable label name."""
    label_names = {
        'wik': 'Wesentliches in Kürze (Cyan)',
        'summary': 'Summary (Cyan)',
        'introduction': 'Auftrag und Vorgehen (Yellow)',
        'findings': 'Findings/Feststellungen (Green)',
        'evaluation': 'Beurteilung der EFK (Blue)',
        'recommendations': 'Empfehlungen (Dark Orange)',
        'empfehlung': 'Empfehlung (Dark Orange)',
        'response': 'Stellungnahme (Pink)',
        'appendix': 'Anhang/Annex (Gray)',
        'background': 'Hintergrund (Sky Blue)',
        'theme': 'Themenbericht (Light Pink)',
        'core_issue': 'Kernproblem (Tomato)',
        'info_a': 'Umwelt, Info A (Pale Green)',
        'info_b': 'Umwelt, Info B (Light Yellow)',
        'info_c': 'Umwelt, Info C (Lavender)',
        'flank_a': 'Flankierend A (Misty Rose)',
        'flank_b': 'Flankierend B (Papaya Whip)',
        'flank_c': 'Flankierend C (Alice Blue)',
        'risk': 'Risiken (Dark Orange)',
        'cost': 'Kosten (Goldenrod)',
        'timeline': 'Zeithorizont (Light Steel Blue)',
        'misc': 'Sonstiges (Light Gray)',
        'unknown': 'Unknown (Gray)'
    }
    return label_names.get(label, f'Unknown ({label})')

def reconstruct_document(chunks_file: Path, output_file: Path):
    """Reconstruct a Word document from chunks with their labels and colors."""
    
    logger.info(f"Loading chunks from: {chunks_file}")
    
    # Load chunks
    with open(chunks_file, 'r', encoding='utf-8') as f:
        chunks = json.load(f)
    
    logger.info(f"Loaded {len(chunks)} chunks")
    
    # Create Word document
    doc = Document()
    
    # Add title
    title = doc.add_heading(chunks_file.stem.replace('_chunks', ''), 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add summary of labels found
    labels_summary = {}
    for chunk in chunks:
        label = chunk.get('label', 'default')
        labels_summary[label] = labels_summary.get(label, 0) + 1
    
    doc.add_heading('Labeling Summary', level=1)
    summary_para = doc.add_paragraph()
    summary_para.add_run(f"Total chunks: {len(chunks)}\n")
    for label, count in sorted(labels_summary.items()):
        run = summary_para.add_run(f"  • {get_label_name(label)}: {count} chunks\n")
        if label in COLOR_MAP:
            run.font.color.rgb = COLOR_MAP[label]
    
    doc.add_page_break()
    
    # Add chunks with their labels and colors
    doc.add_heading('Document Chunks with Labels', level=1)
    
    for i, chunk in enumerate(chunks):
        # Extract chunk info
        chunk_id = chunk.get('chunk_id', f'chunk_{i}')
        label = chunk.get('label', 'default')
        text = chunk.get('text', '')
        start_page = chunk.get('start_page', 'N/A')
        end_page = chunk.get('end_page', 'N/A')
        
        # Add chunk header
        header = doc.add_paragraph()
        header_run = header.add_run(f"Chunk {i+1}: {chunk_id}")
        header_run.bold = True
        header_run.font.size = Pt(12)
        
        # Add metadata
        meta = doc.add_paragraph()
        meta_run = meta.add_run(f"Label: {get_label_name(label)} | Pages: {start_page}-{end_page}")
        meta_run.font.size = Pt(10)
        meta_run.italic = True
        
        # Add chunk text with background color
        chunk_para = doc.add_paragraph()
        chunk_run = chunk_para.add_run(text[:500] + '...' if len(text) > 500 else text)
        chunk_run.font.size = Pt(11)
        
        # Apply color based on label
        if label in COLOR_MAP:
            # Note: python-docx doesn't support background colors directly,
            # so we'll use font color to indicate the label
            chunk_run.font.color.rgb = COLOR_MAP[label]
            
            # Add a colored indicator
            color_indicator = doc.add_paragraph()
            indicator_run = color_indicator.add_run(f"[COLOR: {get_label_name(label)}]")
            indicator_run.font.color.rgb = COLOR_MAP[label]
            indicator_run.font.size = Pt(9)
            indicator_run.bold = True
        
        # Add separator
        doc.add_paragraph('─' * 80)
    
    # Add legend at the end
    doc.add_page_break()
    doc.add_heading('Color Legend', level=1)
    
    for label, color in sorted(COLOR_MAP.items()):
        if label != 'default':
            para = doc.add_paragraph()
            run = para.add_run(f"■ {get_label_name(label)}")
            run.font.color.rgb = color
            run.font.size = Pt(12)
    
    # Save document
    doc.save(output_file)
    logger.info(f"Document saved to: {output_file}")
    
    # Print statistics
    print("\n📊 Labeling Statistics:")
    print(f"Total chunks: {len(chunks)}")
    print("\nLabel distribution:")
    for label, count in sorted(labels_summary.items()):
        percentage = (count / len(chunks)) * 100
        print(f"  {get_label_name(label):30} {count:4} chunks ({percentage:5.1f}%)")

def main():
    parser = argparse.ArgumentParser(description='Reconstruct colored document from chunks')
    parser.add_argument('--doc-id', default='22723BE Prüfbericht V04 - DOC - Prüfung der Investitionsplanung  - farbcoded',
                       help='Document ID to process')
    parser.add_argument('--output-dir', default='08_chunk_visualization',
                       help='Output directory for Word document')
    
    args = parser.parse_args()
    
    # Setup paths
    chunks_file = Path('02_chunked_data') / f"{args.doc_id}_chunks.json"
    output_dir = Path(args.output_dir)
    output_dir.mkdir(exist_ok=True)
    output_file = output_dir / f"{args.doc_id}_reconstructed.docx"
    
    if not chunks_file.exists():
        logger.error(f"Chunks file not found: {chunks_file}")
        return
    
    # Reconstruct document
    reconstruct_document(chunks_file, output_file)

if __name__ == "__main__":
    main()
