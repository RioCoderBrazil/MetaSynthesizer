#!/usr/bin/env python3
"""
Reconstruct Word documents from the properly labeled chunks, including ALL colors.
This includes the previously missing GREEN (findings) and GRAY (appendix) sections.
"""

import json
from pathlib import Path
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_COLOR_INDEX
from collections import defaultdict
import sys

def reconstruct_document(chunks_file: Path, output_file: Path):
    """Reconstruct a Word document from labeled chunks with proper coloring"""
    
    # Load chunks
    with open(chunks_file, 'r', encoding='utf-8') as f:
        chunks = json.load(f)
    
    # Create new document
    doc = Document()
    
    # Add title
    title = doc.add_heading(chunks_file.stem.replace('_highlighted_chunks', ''), 0)
    
    # Color mapping for highlights
    label_to_highlight = {
        'findings': WD_COLOR_INDEX.BRIGHT_GREEN,      # GREEN - Main findings
        'appendix': WD_COLOR_INDEX.GRAY_25,           # GRAY - Annex sections
        'wik': WD_COLOR_INDEX.TURQUOISE,              # CYAN - Executive summary
        'introduction': WD_COLOR_INDEX.YELLOW,         # YELLOW - Introduction
        'evaluation': WD_COLOR_INDEX.BLUE,             # BLUE - Evaluation
        'response': WD_COLOR_INDEX.PINK,               # PINK - Response
        'recommendations': WD_COLOR_INDEX.DARK_YELLOW, # DARK YELLOW - Recommendations
    }
    
    # Label descriptions
    label_descriptions = {
        'findings': 'Feststellungen (Findings)',
        'appendix': 'Anhang (Appendix)',
        'wik': 'Wesentliches in Kürze (Executive Summary)',
        'introduction': 'Auftrag und Vorgehen (Introduction)',
        'evaluation': 'Beurteilung der EFK (Evaluation)',
        'response': 'Stellungnahme (Response)',
        'recommendations': 'Empfehlungen (Recommendations)'
    }
    
    # Add color legend
    doc.add_heading('Color Legend / Farblegende', 1)
    legend_para = doc.add_paragraph()
    for label, description in label_descriptions.items():
        if any(chunk['label'] == label for chunk in chunks):
            run = legend_para.add_run(f'■ {description}\n')
            run.font.highlight_color = label_to_highlight.get(label, WD_COLOR_INDEX.AUTO)
            run.font.size = Pt(12)
    
    # Statistics
    label_stats = defaultdict(int)
    for chunk in chunks:
        label_stats[chunk['label']] += len(chunk['text'])
    
    # Add statistics
    doc.add_heading('Document Statistics / Dokumentstatistik', 1)
    stats_para = doc.add_paragraph()
    
    total_chars = sum(label_stats.values())
    for label, char_count in sorted(label_stats.items(), key=lambda x: x[1], reverse=True):
        percentage = (char_count / total_chars * 100) if total_chars > 0 else 0
        description = label_descriptions.get(label, label)
        stats_para.add_run(f'{description}: {char_count:,} characters ({percentage:.1f}%)\n')
    
    # Special note about findings
    if label_stats.get('findings', 0) > 0:
        findings_pct = (label_stats['findings'] / total_chars * 100)
        note_para = doc.add_paragraph()
        note_run = note_para.add_run(f'\n✓ GREEN FINDINGS INCLUDED: {label_stats["findings"]:,} characters ({findings_pct:.1f}% of document)')
        note_run.font.bold = True
        note_run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
    
    if label_stats.get('appendix', 0) > 0:
        appendix_pct = (label_stats['appendix'] / total_chars * 100)
        note_para = doc.add_paragraph()
        note_run = note_para.add_run(f'✓ GRAY APPENDIX INCLUDED: {label_stats["appendix"]:,} characters ({appendix_pct:.1f}% of document)')
        note_run.font.bold = True
        note_run.font.highlight_color = WD_COLOR_INDEX.GRAY_25
    
    # Add content sections
    doc.add_page_break()
    doc.add_heading('Document Content / Dokumentinhalt', 1)
    
    # Group chunks by label and page
    sections_by_label = defaultdict(list)
    for chunk in chunks:
        sections_by_label[chunk['label']].append(chunk)
    
    # Add sections in order
    section_order = ['wik', 'introduction', 'findings', 'evaluation', 'recommendations', 'response', 'appendix']
    
    for label in section_order:
        if label in sections_by_label:
            # Add section heading
            doc.add_heading(label_descriptions.get(label, label), 2)
            
            # Add chunks for this section
            for chunk in sections_by_label[label]:
                para = doc.add_paragraph()
                
                # Add the text with appropriate highlight
                run = para.add_run(chunk['text'])
                run.font.highlight_color = label_to_highlight.get(label, WD_COLOR_INDEX.AUTO)
                
                # Add page reference
                if 'start_page' in chunk:
                    page_run = para.add_run(f" (S. {chunk['start_page']})")
                    page_run.font.size = Pt(9)
                    page_run.font.italic = True
    
    # Save document
    output_file.parent.mkdir(exist_ok=True)
    doc.save(str(output_file))
    
    print(f"✓ Reconstructed document saved to: {output_file}")
    
    # Print summary
    print("\nDocument Summary:")
    print("-" * 50)
    for label, char_count in sorted(label_stats.items(), key=lambda x: x[1], reverse=True):
        percentage = (char_count / total_chars * 100) if total_chars > 0 else 0
        print(f"  {label_descriptions.get(label, label):40}: {percentage:5.1f}%")

def reconstruct_all_documents():
    """Reconstruct all documents from highlighted chunks"""
    
    chunks_dir = Path("02_chunked_data_highlighted")
    output_dir = Path("08_chunk_visualization_complete")
    output_dir.mkdir(exist_ok=True)
    
    # Get all chunk files
    chunk_files = list(chunks_dir.glob("*_highlighted_chunks.json"))
    
    print(f"Found {len(chunk_files)} documents to reconstruct")
    print("=" * 80)
    
    for chunk_file in chunk_files[:5]:  # Process first 5 as examples
        output_file = output_dir / f"{chunk_file.stem.replace('_highlighted_chunks', '')}_complete_colored.docx"
        reconstruct_document(chunk_file, output_file)
        print()

if __name__ == "__main__":
    # First reconstruct specific document user was interested in
    specific_doc = Path("02_chunked_data_highlighted/22723BE Prüfbericht V04 - DOC - Prüfung der Investitionsplanung  - farbcoded_highlighted_chunks.json")
    if specific_doc.exists():
        print("Reconstructing the document you were working with:")
        print("="*80)
        output = Path("08_chunk_visualization_complete/22723BE Prüfbericht V04 - DOC - Prüfung der Investitionsplanung  - farbcoded_complete_colored.docx")
        reconstruct_document(specific_doc, output)
        print("\n" + "="*80 + "\n")
    
    # Then do others
    reconstruct_all_documents()
